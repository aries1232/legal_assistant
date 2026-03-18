import os
import io
import uuid
from typing import Any, Dict, List

import fitz  # PyMuPDF
import docx
from llama_index.core import Document, StorageContext, VectorStoreIndex
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore

try:
    from src.chroma_client import get_chroma_client
except ModuleNotFoundError:
    from chroma_client import get_chroma_client

EMBED_MODEL = os.getenv("EMBED_MODEL", "BAAI/bge-base-en-v1.5")
CHROMA_COLLECTION = os.getenv("CHROMA_COLLECTION", "legal_documents")
CHUNK_SIZE_TOKENS = int(os.getenv("CHUNK_SIZE_TOKENS", "512"))
CHUNK_OVERLAP_TOKENS = int(os.getenv("CHUNK_OVERLAP_TOKENS", "50"))


class DocumentIngestor:
    def __init__(self) -> None:
        self.db = None
        self.collection = None
        self.vector_store: ChromaVectorStore | None = None
        self.embed_model: HuggingFaceEmbedding | None = None
        self.converter: DocumentConverter | None = None
        self.parser = SentenceSplitter(
            chunk_size=CHUNK_SIZE_TOKENS,
            chunk_overlap=CHUNK_OVERLAP_TOKENS,
        )

    def _ensure_ready(self) -> None:
        if self.db is None:
            self.db = get_chroma_client()
        if self.collection is None:
            self.collection = self.db.get_or_create_collection(CHROMA_COLLECTION)
        if self.vector_store is None:
            self.vector_store = ChromaVectorStore(chroma_collection=self.collection)
        if self.embed_model is None:
            self.embed_model = HuggingFaceEmbedding(model_name=EMBED_MODEL)

    def _extract_text(self, filename: str, file_bytes: bytes) -> str:
        self._ensure_ready()
        ext = os.path.splitext(filename)[1].lower()
        extracted_text = ""

        try:
            if ext == ".pdf":
                # Use PyMuPDF for fast local PDF extraction
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                extracted_text = "\n".join([page.get_text() for page in doc])
                doc.close()
            elif ext == ".docx":
                # Use python-docx for DOCX files
                doc = docx.Document(io.BytesIO(file_bytes))
                extracted_text = "\n".join([para.text for para in doc.paragraphs])
            elif ext in [".txt", ".md", ".csv"]:
                # Native fast text decoding
                extracted_text = file_bytes.decode("utf-8", errors="replace")
            else:
                raise ValueError(f"Unsupported file format: {ext}")
        except Exception as e:
            print(f"Error extracting text from {filename}: {e}")
            
        return extracted_text

    def ingest_document(self, filename: str, file_bytes: bytes, doc_id: str, status_callback=None) -> Dict[str, Any]:
        try:
            if status_callback: status_callback(f"Extracting text from {filename}...", 20)
            extracted_text = self._extract_text(filename, file_bytes)
            if not extracted_text.strip():
                return {
                    "ok": False,
                    "doc_id": doc_id,
                    "filename": filename,
                    "chunk_count": 0,
                    "error": "No extractable text found in document.",
                }

            if status_callback: status_callback("Parsing and splitting text...", 50)
            doc = Document(
                text=extracted_text,
                metadata={
                    "doc_id": doc_id,
                    "filename": filename,
                    "source_type": os.path.splitext(filename)[1].lower().strip('.'),
                },
            )
            nodes = self.parser.get_nodes_from_documents([doc])
            
            if status_callback: status_callback(f"Generating embeddings for {len(nodes)} chunks...", 75)
            self._ensure_ready()
            storage_context = StorageContext.from_defaults(vector_store=self.vector_store)
            index = VectorStoreIndex(
                nodes,
                storage_context=storage_context,
                embed_model=self.embed_model,
                show_progress=False,
            )
            
            if status_callback: status_callback("Finalizing index...", 100)
            return {
                "ok": True,
                "doc_id": doc_id,
                "filename": filename,
                "chunk_count": len(nodes),
                "error": None,
            }
        except Exception as e:
            error_message = str(e)
            if "expecting embedding with dimension of 768, got 384" in error_message.lower():
                error_message = (
                    "Embedding model mismatch: the Chroma collection expects 768-dimensional vectors, "
                    f"but the configured model `{EMBED_MODEL}` produced 384-dimensional vectors. "
                    "Use `BAAI/bge-base-en-v1.5` for both ingestion and querying, or recreate the collection "
                    "if you intentionally want to switch models."
                )
            return {
                "ok": False,
                "doc_id": doc_id,
                "filename": filename,
                "chunk_count": 0,
                "error": error_message,
            }


def generate_doc_id() -> str:
    return str(uuid.uuid4())
